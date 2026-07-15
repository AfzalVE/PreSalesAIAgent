import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding margins (in dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_proposal_docx(proposal_data: dict, output_filepath: str):
    """
    Generates a professional DOCX document based on proposal_data
    following the format and style of City_Canvas_POC.docx.
    """
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51) # Charcoal

    # Title Page / Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title_tag = p_title.add_run("PROOF OF CONCEPT (POC) & DEVELOPMENT PROPOSAL\n")
    run_title_tag.font.size = Pt(12)
    run_title_tag.font.bold = True
    run_title_tag.font.color.rgb = RGBColor(12, 60, 96) # Dark Navy

    run_title = p_title.add_run(proposal_data.get("project_name", "AI System Development").upper() + "\n")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(12, 60, 96)

    run_subtitle = p_title.add_run("Custom Software Engineering and Implementation Blueprint\n")
    run_subtitle.font.size = Pt(14)
    run_subtitle.font.italic = True
    run_subtitle.font.color.rgb = RGBColor(128, 128, 128)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run(f"Document Version: 1.0\n")
    p_meta.add_run(f"Estimated Budget: ${proposal_data.get('estimated_cost', 0.0):,.2f} USD\n")
    p_meta.add_run(f"Development Timeline: {proposal_data.get('estimated_duration', '12 Weeks')}\n")
    p_meta.add_run(f"Proposal Type: {proposal_data.get('proposal_type', 'FULL')}\n")

    # Add spacing
    doc.add_paragraph()

    # Helper function to add styled headings
    def add_styled_heading(text, level=1):
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(12, 60, 96)
            # Add subtle bottom border style spacing
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(8)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(12, 60, 96)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(4)
        return heading

    # 1. Executive Summary
    add_styled_heading("1. Executive Summary")
    doc.add_paragraph(
        f"This development proposal outlines the technical specifications, scope, resource planning, and timeline milestones "
        f"for the '{proposal_data.get('project_name')}' project. Our approach balances cutting-edge engineering with "
        f"speed-to-market methodologies to deliver a premium, robust digital solution."
    )

    # 2. Scope & Proposed Solution
    add_styled_heading("2. Proposed Solution & Scope")
    doc.add_paragraph(proposal_data.get("scope", "Detailed implementation scope and architectural integration plans."))

    # 3. Technology Stack
    add_styled_heading("3. Proposed Technology Stack")
    tech_stack = proposal_data.get("tech_stack") or {}
    
    # Create Table 2: Tech Stack
    table_tech = doc.add_table(rows=1, cols=3)
    table_tech.style = 'Table Grid'
    hdr_cells = table_tech.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology / Framework'
    hdr_cells[2].text = 'Purpose'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0C3C60")
        set_cell_margins(cell)

    for layer, tech in tech_stack.items():
        row_cells = table_tech.add_row().cells
        row_cells[0].text = str(layer).capitalize()
        row_cells[1].text = str(tech)
        row_cells[2].text = f"Primary platform layer for the {layer} requirements."
        for cell in row_cells:
            set_cell_margins(cell)

    doc.add_paragraph() # Spacer

    # 4. Project Development Plan & Timeline Milestones
    add_styled_heading("4. Development Plan & Timeline")
    doc.add_paragraph(
        f"The development cycle is projected at {proposal_data.get('estimated_duration')}. "
        f"Following the structure of our standard Proof of Concept delivery pipeline, the specific deliverables and timeline phases are detailed below:"
    )

    # Create Table 5: Development Plan
    timeline_phases = proposal_data.get("timeline_phases") or []
    if timeline_phases:
        table_plan = doc.add_table(rows=1, cols=3)
        table_plan.style = 'Table Grid'
        hdr_cells = table_plan.rows[0].cells
        hdr_cells[0].text = 'Phase'
        hdr_cells[1].text = 'Duration'
        hdr_cells[2].text = 'Output'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, "0C3C60")
            set_cell_margins(cell)

        for phase in timeline_phases:
            row_cells = table_plan.add_row().cells
            row_cells[0].text = phase.get("Phase", "N/A")
            row_cells[1].text = phase.get("Duration", "N/A")
            row_cells[2].text = phase.get("Output", "N/A")
            for cell in row_cells:
                set_cell_margins(cell)

    doc.add_paragraph() # Spacer

    # 5. Assumptions & Key Risks
    add_styled_heading("5. Assumptions & Key Risks")
    
    add_styled_heading("Assumptions", level=2)
    doc.add_paragraph(proposal_data.get("assumptions") or "No technical assumptions listed.")

    add_styled_heading("Risks & Mitigation", level=2)
    doc.add_paragraph(proposal_data.get("risks") or "No project risks listed.")

    # 6. Selected Resources Allocation
    add_styled_heading("6. Allocated Technical Resources")
    selected_resources = proposal_data.get("selected_resources", {}).get("resources", [])
    
    if selected_resources:
        table_res = doc.add_table(rows=1, cols=3)
        table_res.style = 'Table Grid'
        hdr_cells = table_res.rows[0].cells
        hdr_cells[0].text = 'Resource Name'
        hdr_cells[1].text = 'Assigned Role'
        hdr_cells[2].text = 'Allocated Hours'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, "0C3C60")
            set_cell_margins(cell)

        for res in selected_resources:
            row_cells = table_res.add_row().cells
            row_cells[0].text = res.get("name", "N/A")
            row_cells[1].text = res.get("role", "N/A")
            row_cells[2].text = f"{res.get('allocated_hours', 0)} Hours"
            for cell in row_cells:
                set_cell_margins(cell)

    doc.add_paragraph() # Spacer

    # Save document
    os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
    doc.save(output_filepath)
    print(f"[DOCX] Final proposal Word document generated at {output_filepath}")

