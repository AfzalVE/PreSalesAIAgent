"""
document_builder.py
===================
Converts rendered HTML into PDF (via WeasyPrint) and DOCX (via python-docx).

PDF:
    WeasyPrint parses the self-contained HTML + CSS and produces a high-fidelity PDF.

DOCX:
    python-docx is used to produce a structured Word document with professional
    formatting. Note: DOCX output is a structural representation — full pixel-perfect
    CSS fidelity is not achievable in DOCX format.

Both functions return raw bytes so the caller can stream them over HTTP,
write them to disk, or attach them to emails.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# PDF Generation — WeasyPrint
# ---------------------------------------------------------------------------


def generate_pdf(html_content: str) -> bytes:
    """
    Convert a self-contained HTML string to PDF bytes using WeasyPrint.

    Args:
        html_content: The full HTML document string (from renderer.render_html).

    Returns:
        Raw PDF bytes.

    Raises:
        ImportError: If WeasyPrint is not installed.
        Exception:   For any WeasyPrint rendering error.
    """
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
    except ImportError as exc:
        raise ImportError(
            "WeasyPrint is required for PDF generation. "
            "Install it with: pip install weasyprint"
        ) from exc

    logger.info("Generating PDF via WeasyPrint...")

    font_config = FontConfiguration()

    # Additional print CSS to ensure good page breaks and margins
    print_css = CSS(
        string="""
        @page {
            margin: 0;
            size: A4;
        }
        body {
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }
        """,
        font_config=font_config,
    )

    try:
        pdf_bytes = (
            HTML(string=html_content, base_url=None)
            .write_pdf(stylesheets=[print_css], font_config=font_config)
        )
        logger.info("PDF generated | size=%d bytes", len(pdf_bytes))
        return pdf_bytes
    except Exception as exc:
        logger.error("WeasyPrint PDF generation failed: %s", exc)
        raise


# ---------------------------------------------------------------------------
# DOCX Generation — python-docx
# ---------------------------------------------------------------------------


def generate_docx(proposal_output, proposal_id: str) -> bytes:
    """
    Build a professionally formatted Word (.docx) document from a ProposalOutput.

    Args:
        proposal_output: The ProposalOutput Pydantic model.
        proposal_id:     Unique proposal reference string.

    Returns:
        Raw DOCX bytes.

    Raises:
        ImportError: If python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install it with: pip install python-docx"
        ) from exc

    logger.info("Generating DOCX via python-docx...")

    # Color palette (RGB tuples)
    PRIMARY      = RGBColor(0x1A, 0x2E, 0x6F)
    ACCENT       = RGBColor(0x3D, 0x7B, 0xFF)
    GOLD         = RGBColor(0xF5, 0xA6, 0x23)
    TEXT_LIGHT   = RGBColor(0x71, 0x80, 0x96)
    WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
    BG_LIGHT_RGB = RGBColor(0xF7, 0xF9, 0xFC)

    p = proposal_output
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ----------------------------------------------------------------
    # Utility helpers
    # ----------------------------------------------------------------

    def add_heading(text: str, level: int = 1, color: RGBColor = PRIMARY) -> None:
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = color
            run.font.bold = True

    def add_subheading(text: str, color: RGBColor = ACCENT) -> None:
        p_obj = doc.add_paragraph()
        run = p_obj.add_run(text.upper())
        run.font.color.rgb = color
        run.font.size = Pt(8)
        run.font.bold = True

    def add_para(text: str, italic: bool = False, color: RGBColor = None) -> None:
        if not text.strip():
            return
        para = doc.add_paragraph()
        run = para.add_run(text.strip())
        run.font.size = Pt(10.5)
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color

    def add_bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_separator() -> None:
        doc.add_paragraph("─" * 80).runs[0].font.color.rgb = TEXT_LIGHT

    def shade_row(row, color_hex: str = "EBF1FF") -> None:
        for cell in row.cells:
            tc = cell._tc
            tcp = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), color_hex)
            tcp.append(shd)

    # ================================================================
    # COVER PAGE
    # ================================================================
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title block
    title_run = cover_para.add_run(f"\n\n{p.project.title}\n")
    title_run.font.size      = Pt(28)
    title_run.font.bold      = True
    title_run.font.color.rgb = PRIMARY

    sub_run = cover_para.add_run(f"{p.project.description}\n\n")
    sub_run.font.size      = Pt(12)
    sub_run.font.color.rgb = TEXT_LIGHT
    sub_run.font.italic    = True

    # Meta info box as a table
    meta_table = doc.add_table(rows=2, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For",  f"{p.client.contact}\n{p.client.company}"),
        ("Investment",    f"{p.raw_budget.currency} {p.raw_budget.amount:,.0f}"),
        ("Duration",      p.timeline.duration),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_table.rows[0].cells[i]
        value_cell = meta_table.rows[1].cells[i]
        label_cell.text = label
        value_cell.text = value
        for cell, fsize, bold, color in [
            (label_cell, 8, True, TEXT_LIGHT),
            (value_cell, 11, True, PRIMARY),
        ]:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(fsize)
                    run.font.bold = bold
                    run.font.color.rgb = color
        shade_row(meta_table.rows[0], "1A2E6F")
        shade_row(meta_table.rows[1], "EBF1FF")

    doc.add_paragraph(f"\nProposal ID: {proposal_id}").runs[0].font.color.rgb = TEXT_LIGHT
    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    add_heading("Table of Contents", level=1)
    toc_items = [
        "01. Executive Summary",
        "02. Proposed Solution",
        "03. Technical Architecture",
        "04. Project Timeline",
        "05. Investment & Budget",
        "06. Risks & Mitigation",
        "07. Project Deliverables",
        "08. Proposed Team",
    ]
    for item in toc_items:
        add_bullet(item)
    doc.add_page_break()

    # ================================================================
    # 01. EXECUTIVE SUMMARY
    # ================================================================
    add_subheading("Section 01")
    add_heading("Executive Summary", level=1)

    # Headline box
    hl_para = doc.add_paragraph()
    hl_run = hl_para.add_run(f'"{p.executive_summary.headline}"')
    hl_run.font.size = Pt(13)
    hl_run.font.bold = True
    hl_run.font.color.rgb = PRIMARY
    hl_run.font.italic = True
    shade_row(hl_para._p, "EBF1FF")

    for para in p.executive_summary.body.split("\n\n"):
        add_para(para)

    add_separator()

    # ================================================================
    # 02. PROPOSED SOLUTION
    # ================================================================
    add_subheading("Section 02")
    add_heading("Proposed Solution", level=1)

    for para in p.solution_description.overview.split("\n\n"):
        add_para(para)

    doc.add_paragraph("\nKey Business Benefits:").runs[0].font.bold = True
    for benefit in p.solution_description.key_benefits:
        add_bullet(benefit)

    doc.add_paragraph("\nFeature Highlights:").runs[0].font.bold = True
    for feat in p.solution_description.feature_highlights:
        add_bullet(feat)

    add_separator()

    # ================================================================
    # 03. TECHNICAL ARCHITECTURE
    # ================================================================
    add_subheading("Section 03")
    add_heading("Technical Architecture", level=1)

    # Tech stack table
    ts = p.tech_stack
    stack_items = [
        ("Frontend", ts.frontend),
        ("Backend", ts.backend),
        ("Database", ts.database),
        ("Hosting / Cloud", ts.hosting),
    ]
    stack_rows = [(l, v) for l, v in stack_items if v]

    if stack_rows:
        tbl = doc.add_table(rows=1 + len(stack_rows), cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0]
        hdr.cells[0].text = "Layer"
        hdr.cells[1].text = "Technology"
        shade_row(hdr, "1A2E6F")
        for cell in hdr.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold  = True
                    run.font.color.rgb = WHITE
        for i, (label, value) in enumerate(stack_rows, start=1):
            row = tbl.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value or "—"
            if i % 2 == 0:
                shade_row(row, "F7F9FC")

    doc.add_paragraph()

    doc.add_heading("Architecture Overview", level=2)
    for para in p.technical_details.architecture_overview.split("\n\n"):
        add_para(para)

    doc.add_heading("Technology Selection Rationale", level=2)
    for para in p.technical_details.stack_rationale.split("\n\n"):
        add_para(para)

    doc.add_heading("Integration & Security Notes", level=2)
    for para in p.technical_details.integration_notes.split("\n\n"):
        add_para(para)

    add_separator()

    # ================================================================
    # 04. PROJECT TIMELINE
    # ================================================================
    add_subheading("Section 04")
    add_heading("Project Timeline", level=1)

    for para in p.timeline.overview.split("\n\n"):
        add_para(para)

    doc.add_paragraph()
    phase_tbl = doc.add_table(rows=1 + len(p.timeline.phases), cols=3)
    phase_tbl.style = "Table Grid"
    hdrs = phase_tbl.rows[0]
    hdrs.cells[0].text = "Phase"
    hdrs.cells[1].text = "Duration"
    hdrs.cells[2].text = "Description"
    shade_row(hdrs, "1A2E6F")
    for cell in hdrs.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE

    for i, phase in enumerate(p.timeline.phases, start=1):
        row = phase_tbl.rows[i]
        row.cells[0].text = phase.get("name", "")
        row.cells[1].text = f"{phase.get('weeks', '')} weeks"
        row.cells[2].text = phase.get("description", "")
        if i % 2 == 0:
            shade_row(row, "F7F9FC")

    doc.add_paragraph("\nKey Milestones:").runs[0].font.bold = True
    for ms in p.timeline.milestones:
        add_bullet(ms)

    add_separator()

    # ================================================================
    # 05. INVESTMENT & BUDGET
    # ================================================================
    add_subheading("Section 05")
    add_heading("Investment & Budget", level=1)

    # Total callout
    total_para = doc.add_paragraph()
    t_label = total_para.add_run("Total Investment: ")
    t_label.font.size = Pt(12)
    t_label.font.bold = True
    t_label.font.color.rgb = PRIMARY
    t_val = total_para.add_run(f"{p.raw_budget.currency} {p.raw_budget.amount:,.0f}")
    t_val.font.size = Pt(16)
    t_val.font.bold = True
    t_val.font.color.rgb = GOLD

    for para in p.budget_section.summary.split("\n\n"):
        add_para(para)

    doc.add_paragraph()
    budget_rows = p.budget_section.breakdown
    b_tbl = doc.add_table(rows=1 + len(budget_rows) + 1, cols=2)
    b_tbl.style = "Table Grid"
    b_hdr = b_tbl.rows[0]
    b_hdr.cells[0].text = "Line Item"
    b_hdr.cells[1].text = f"Cost ({p.raw_budget.currency})"
    shade_row(b_hdr, "1A2E6F")
    for cell in b_hdr.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE

    for i, row_data in enumerate(budget_rows, start=1):
        row = b_tbl.rows[i]
        row.cells[0].text = row_data.get("item", "")
        row.cells[1].text = f"{row_data.get('cost', 0):,.0f}"
        if i % 2 == 0:
            shade_row(row, "F7F9FC")

    # Total row
    total_row = b_tbl.rows[-1]
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = f"{p.raw_budget.amount:,.0f}"
    shade_row(total_row, "EBF1FF")
    for cell in total_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = PRIMARY

    doc.add_paragraph()
    terms_para = doc.add_paragraph()
    terms_label = terms_para.add_run("Payment Terms: ")
    terms_label.font.bold = True
    terms_label.font.color.rgb = PRIMARY
    terms_para.add_run(p.budget_section.payment_terms)

    add_separator()

    # ================================================================
    # 06. RISKS & MITIGATION
    # ================================================================
    add_subheading("Section 06")
    add_heading("Risks & Mitigation", level=1)

    for para in p.risks.intro.split("\n\n"):
        add_para(para)

    doc.add_paragraph()
    risk_rows = p.risks.risks
    r_tbl = doc.add_table(rows=1 + len(risk_rows), cols=3)
    r_tbl.style = "Table Grid"
    r_hdr = r_tbl.rows[0]
    r_hdr.cells[0].text = "Risk"
    r_hdr.cells[1].text = "Potential Impact"
    r_hdr.cells[2].text = "Mitigation Strategy"
    shade_row(r_hdr, "1A2E6F")
    for cell in r_hdr.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE

    for i, risk in enumerate(risk_rows, start=1):
        row = r_tbl.rows[i]
        row.cells[0].text = risk.get("risk", "")
        row.cells[1].text = risk.get("impact", "")
        row.cells[2].text = risk.get("mitigation", "")
        if i % 2 == 0:
            shade_row(row, "F7F9FC")

    add_separator()

    # ================================================================
    # 07. DELIVERABLES
    # ================================================================
    add_subheading("Section 07")
    add_heading("Project Deliverables", level=1)

    for para in p.deliverables.intro.split("\n\n"):
        add_para(para)

    doc.add_paragraph()
    for d in p.deliverables.deliverables:
        add_bullet(f"✓  {d}")

    add_separator()

    # ================================================================
    # 08. PROPOSED TEAM
    # ================================================================
    add_subheading("Section 08")
    add_heading("Proposed Team", level=1)

    add_para(
        f"Our carefully selected team brings deep domain expertise and a proven track record "
        f"of delivering complex software projects on time and within budget."
    )

    doc.add_paragraph()
    team_tbl = doc.add_table(rows=1 + len(p.team), cols=2)
    team_tbl.style = "Table Grid"
    t_hdr = team_tbl.rows[0]
    t_hdr.cells[0].text = "Role"
    t_hdr.cells[1].text = "Experience"
    shade_row(t_hdr, "1A2E6F")
    for cell in t_hdr.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE

    for i, member in enumerate(p.team, start=1):
        row = team_tbl.rows[i]
        row.cells[0].text = member.role
        row.cells[1].text = member.experience
        if i % 2 == 0:
            shade_row(row, "F7F9FC")

    if p.assumptions:
        doc.add_paragraph("\nProject Assumptions:").runs[0].font.bold = True
        for assumption in p.assumptions:
            add_bullet(assumption)

    # ================================================================
    # CLOSING
    # ================================================================
    doc.add_page_break()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = closing.add_run(f"\n\nReady to Build Together?\n\n")
    c_run.font.size      = Pt(22)
    c_run.font.bold      = True
    c_run.font.color.rgb = PRIMARY

    sub_run = closing.add_run(
        f"We are excited about the opportunity to partner with {p.client.company}.\n"
        f"Proposal Reference: {proposal_id}\n\n"
        f"Confidential — Prepared exclusively for {p.client.company}"
    )
    sub_run.font.size      = Pt(11)
    sub_run.font.color.rgb = TEXT_LIGHT

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()

    logger.info("DOCX generated | size=%d bytes", len(docx_bytes))
    return docx_bytes
